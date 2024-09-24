import streamlit as st
from datetime import datetime, date
from PIL import Image as PILImage
import numpy as np
# import io
from docx import Document
from docx.shared import Inches
from streamlit_drawable_canvas import st_canvas
import smtplib
from email.message import EmailMessage
import shutil
import re
import time
import requests

# Set page configuration with a favicon
st.set_page_config(
    page_title="Multiply Surrey",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"  # "centered" or "wide"
)

# Initialize session state
if 'step' not in st.session_state:
    st.session_state.step = 1
    st.session_state.submission_done = False
    # Section 1a: Engagement Session Details (less than 2hrs)
    if 'course_title_1a' not in st.session_state: st.session_state.course_title_1a = ""
    if 'delivery_location_1a' not in st.session_state: st.session_state.delivery_location_1a = ""
    if 'start_date_1a' not in st.session_state: st.session_state.start_date_1a = None
    if 'end_date_1a' not in st.session_state: st.session_state.end_date_1a = None
    if 'num_hours_1a' not in st.session_state: st.session_state.num_hours_1a = 0
    

    # Section 1b: Substantive Numeracy Delivery Course Details (more than 2 hrs)
    if 'course_code_1b' not in st.session_state: st.session_state.course_code_1b = ""
    if 'course_title_1b' not in st.session_state: st.session_state.course_title_1b = ""
    if 'delivery_location_1b' not in st.session_state: st.session_state.delivery_location_1b = ""
    if 'start_date_1b' not in st.session_state: st.session_state.start_date_1b = None
    if 'end_date_1b' not in st.session_state: st.session_state.end_date_1b = None
    if 'hours_per_week_1b' not in st.session_state: st.session_state.hours_per_week_1b = 0
    if 'total_weeks_1b' not in st.session_state: st.session_state.total_weeks_1b = 0
    if 'total_glh_1b' not in st.session_state: st.session_state.total_glh_1b = 0

    # Section 1c
    if 'course_code_1c' not in st.session_state: st.session_state.course_code_1c = ""
    if 'course_title_1c' not in st.session_state: st.session_state.course_title_1c = ""
    if 'delivery_location_1c' not in st.session_state: st.session_state.delivery_location_1c = ""
    if 'start_date_1c' not in st.session_state: st.session_state.start_date_1c = None
    if 'end_date_1c' not in st.session_state: st.session_state.end_date_1c = None
    if 'hours_per_week_1c' not in st.session_state: st.session_state.hours_per_week_1c = 0
    if 'total_weeks_1c' not in st.session_state: st.session_state.total_weeks_1c = 0
    if 'total_glh_1c' not in st.session_state: st.session_state.total_glh_1c = 0

    # Section 2: Personal Details
    if 'title' not in st.session_state: st.session_state.title = "Mr"
    if 'legal_sex' not in st.session_state: st.session_state.legal_sex = "Male"
    if 'dob' not in st.session_state: st.session_state.dob = None
    if 'ni_number' not in st.session_state: st.session_state.ni_number = ""
    if 'forename' not in st.session_state: st.session_state.forename = ""
    if 'surname' not in st.session_state: st.session_state.surname = ""
    if 'previous_surname' not in st.session_state: st.session_state.previous_surname = ""
    if 'current_address' not in st.session_state: st.session_state.current_address = ""
    if 'town' not in st.session_state: st.session_state.town = ""
    if 'postcode' not in st.session_state: st.session_state.postcode = ""
    if 'previous_postcodes' not in st.session_state: st.session_state.previous_postcodes = []
    if 'telephone' not in st.session_state: st.session_state.telephone = ""
    if 'email' not in st.session_state: st.session_state.email = ""

    # Section 3: LRS Privacy Notice
    if 'maths_qualification' not in st.session_state: st.session_state.maths_qualification = "No"

    # Section 4: Residency Eligibility
    if 'uk_eea_residency' not in st.session_state: st.session_state.uk_eea_residency = "No"
    if 'nationality_status' not in st.session_state: st.session_state.nationality_status = "No"

    # Section 5: Ethnicity
    if 'ethnicity_category' not in st.session_state: st.session_state.ethnicity_category = "Select a category"
    if 'ethnicity_detail' not in st.session_state: st.session_state.ethnicity_detail = "Select an option"

    # Section 6: Learning Difficulties, Disabilities and Health Problems
    if 'has_difficulties' not in st.session_state: st.session_state.has_difficulties = "No"
    if 'selected_difficulties' not in st.session_state: st.session_state.selected_difficulties = []
    if 'most_affecting' not in st.session_state: st.session_state.most_affecting = ""

    # Section 7: Highest Qualification
    if 'qualification' not in st.session_state: st.session_state.qualification = "No qualifications"

    # Section 8: Employment & Benefit
    if 'employment_status' not in st.session_state: st.session_state.employment_status = ""
    if 'working_hours' not in st.session_state: st.session_state.working_hours = ""
    if 'unemployment_duration' not in st.session_state: st.session_state.unemployment_duration = ""
    if 'job_seekers_allowance' not in st.session_state: st.session_state.job_seekers_allowance = False
    if 'esa' not in st.session_state: st.session_state.esa = False
    if 'universal_credit' not in st.session_state: st.session_state.universal_credit = False
    if 'other_benefit' not in st.session_state: st.session_state.other_benefit = False



    # Section 10: CLS Marketing and Permissions
    if 'marketing_courses_offers' not in st.session_state: st.session_state.marketing_courses_offers = False
    if 'research_profiling' not in st.session_state: st.session_state.research_profiling = False
    if 'contact_email' not in st.session_state: st.session_state.contact_email = False
    if 'contact_phone' not in st.session_state: st.session_state.contact_phone = False
    if 'contact_post' not in st.session_state: st.session_state.contact_post = False
    if 'contact_text' not in st.session_state: st.session_state.contact_text = False

    # Section 11: Learner Declaration
    if 'signature' not in st.session_state: st.session_state.signature = None
    if 'signature_date' not in st.session_state: st.session_state.signature_date = None    

def last():
    st.session_state.clear()

def is_valid_email(email):
    # Comprehensive regex for email validation
    pattern = r'''
        ^                         # Start of string
        (?!.*[._%+-]{2})          # No consecutive special characters
        [a-zA-Z0-9._%+-]{1,64}    # Local part: allowed characters and length limit
        (?<![._%+-])              # No special characters at the end of local part
        @                         # "@" symbol
        [a-zA-Z0-9.-]+            # Domain part: allowed characters
        (?<![.-])                 # No special characters at the end of domain
        \.[a-zA-Z]{2,}$           # Top-level domain with minimum 2 characters
    '''
    
    # Match the entire email against the pattern
    return re.match(pattern, email, re.VERBOSE) is not None

# Sanitize the file name to avoid invalid characters
def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '', filename)

# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files=None, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = ", ".join(receiver_email)
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Attach uploaded files
    if files:
        for uploaded_file in files:
            uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
            msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use the SMTP server for sending the email
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)

def is_signature_drawn(signature):
    # Check if signature is None or an empty numpy array
    if signature is None:
        return False
    # Ensure it is a numpy array and has content
    if isinstance(signature, np.ndarray) and signature.size > 0:
        # Additional check: if the array is not just empty white pixels
        # Assuming white background is [255, 255, 255] in RGB
        if np.all(signature == 255):
            return False
        return True
    return False

def resize_image_to_fit_cell(image, max_width, max_height):
    width, height = image.size
    aspect_ratio = width / height

    if width > max_width:
        width = max_width
        height = int(width / aspect_ratio)

    if height > max_height:
        height = max_height
        width = int(height * aspect_ratio)

    return image.resize((width, height))


def replace_placeholders(template_file, modified_file, placeholder_values, resized_image_path):
    try:
        print(f"Copying template file '{template_file}' to '{modified_file}'...")
        shutil.copy(template_file, modified_file)

        print(f"Opening document '{modified_file}'...")
        doc = Document(modified_file)

        # Function to convert value to string, handling datetime.date objects
        def convert_to_str(value):
            if isinstance(value, date):
                return value.strftime('%Y-%m-%d')  # Convert date to string
            return str(value)  # Convert other types to string

        # Compile regular expressions for all placeholders
        placeholders = {re.escape(key): convert_to_str(value) for key, value in placeholder_values.items()}
        placeholders_pattern = re.compile(r'\b(' + '|'.join(placeholders.keys()) + r')\b')

        # Replace placeholders in paragraphs
        print("Replacing placeholders in paragraphs...")
        for para in doc.paragraphs:
            original_text = para.text
            updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
            if original_text != updated_text:
                print(f"Updated paragraph text: '{original_text}' -> '{updated_text}'")
                para.text = updated_text

        # Replace placeholders in tables
        print("Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
                        if original_text != updated_text:
                            print(f"Updated table cell text: '{original_text}' -> '{updated_text}'")
                            para.text = updated_text

                    # Inspect cell runs
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text
                            run_updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], run_text)
                            if run_text != run_updated_text:
                                print(f"Updated run text in table cell: '{run_text}' -> '{run_updated_text}'")
                                run.text = run_updated_text

        # Check and handle signature placeholder
        print("Inspecting document for 'ph_signature' placeholder...")
        signature_placeholder_found = False

        # Check paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()  # Remove any extra spaces around text
            while 'ph_signature' in para_text:
                print(f"Found 'ph_signature' in paragraph: '{para_text}'")
                para_text = para_text.replace('ph_signature', '').strip()  # Remove 'ph_signature' and any leading/trailing spaces
                para.text = para_text
                
                try:
                    # Add picture to the paragraph
                    print(f"Adding picture to paragraph from path: {resized_image_path}")
                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                    print("Inserted signature image into paragraph.")
                    signature_placeholder_found = True
                except Exception as img_e:
                    print(f"An error occurred with image processing: {img_e}")

        # Check table cells again in case the placeholder was missed
        if not signature_placeholder_found:
            print("Checking table cells for 'ph_signature'...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_text = para.text.strip()
                            while 'ph_signature' in para_text:
                                print(f"Found 'ph_signature' in table cell paragraph: '{para_text}'")
                                para_text = para_text.replace('ph_signature', '').strip()
                                para.text = para_text
                                
                                try:
                                    # Add picture to the table cell
                                    print(f"Adding picture to table cell from path: {resized_image_path}")
                                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                                    print("Inserted signature image into table cell.")
                                    signature_placeholder_found = True
                                except Exception as img_e:
                                    print(f"An error occurred with image processing: {img_e}")

        if not signature_placeholder_found:
            print("No signature placeholder found.")

        # Save the modified document
        print(f"Saving modified document '{modified_file}'...")
        doc.save(modified_file)
        print(f"Document modification complete: '{modified_file}'")

    except Exception as e:
        print(f"An error occurred: {e}")

    # file download button
    # with open(modified_file, 'rb') as f:
    #     file_contents = f.read()
    #     st.download_button(
    #         label="Download Your Response",
    #         data=file_contents,
    #         file_name=modified_file,
    #         mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    #     )


if 'files' not in st.session_state:
    st.session_state.files = []

# Define a function to calculate progress and percentage
def get_progress(step, total_steps=14):
    return int((step / total_steps) * 100)




# Define the total number of steps
total_steps = 12
# Calculate the current progress
progress = get_progress(st.session_state.step, total_steps)
# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)




# Define different steps
if st.session_state.step == 1:
    st.image('resources/header.png', use_column_width=True)

    st.title("LEARNING AGREEMENT 2023 / 2024")
    # st.write("Provider: Prevista Ltd. | Sponsor: Surrey County Council | Website: www.prevista.co.uk")
    # st.write("________________________________________")
    # st.write("**Application and Enrollment Form**")

    # Add question with a dropdown menu
    support_options = [
    "    ", 
    "Self Completing",
    "Innovator Recruitment Team",
    "Catalyst Recruitment Team",
    "Guildford JCP",
    "Camberley JCP",
    "Epsom JCP",
    "Woking JCP",
    "Redhill JCP",
    "Staines JCP",
    "Leatherhead JCP",
    "Croydon JCP",
    "Surrey County Council",
    "Surrey Employment and Skills Board (SESB)",
    "Federation of Small Businesses (FSB) Surrey",
    "Surrey Chambers of Commerce",
    "Voluntary Action South West Surrey",
    "Guildford Borough Council",
    "Woking Borough Council",
    "Surrey Choices",
    "Elmbridge Community Job Club",
    "Mole Valley Employment Group",
    "Surrey Lifelong Learning Partnership (SLLP)",
]
    # st.session_state.selected_option = st.selectbox(
    # "Who is supporting you to fill this form?", 
    # support_options
# )
 
    
    hear_about_options = [
    "Self-referral", 
    "Jobcentre Plus (JCP)",
    "Local Council",
    "Online",
    "Word of Mouth",
    "Community Organization",
    "Employer or Training Provider",
    "Promotional Materials",
    "Other (please specify)"
]
#     st.session_state.hear_about = st.selectbox(
#     "Hear about this opportunity:", 
#     hear_about_options
# )
#     # If the user selects "Other (please specify)", display an input field
#     st.session_state.hother_source=''
#     if st.session_state.hear_about == "Other (please specify)":
#         st.session_state.hother_source = st.text_input("Please specify:")


    st.write("""
    All sections MUST be completed in full.
    """)

    # if st.button("Next"):
    #     if (st.session_state.selected_option!='    '):
    #         st.session_state.step = 2
    #         st.experimental_rerun()
    #     else:
    #         st.warning("Please Choose Valid Support Option.")

    if st.button("Next"):
        st.session_state.step = 2
        st.experimental_rerun()

elif st.session_state.step == 2:
    st.title("> Section 1: SESSION | COURSE DETAILS ")

    # Section 1a: Engagement Session Details (less than 2hrs)
    st.subheader("Section 1a: ENGAGEMENT SESSION DETAILS (less than 2hrs)")
    st.session_state.course_title_1a = st.text_input("Course Title (less than 2hrs)", value=st.session_state.get("course_title_1a", ""))
    st.session_state.delivery_location_1a = st.text_input("Delivery Location (less than 2hrs)", value=st.session_state.get("delivery_location_1a", ""))
    
    # Handle Start Date and End Date as date objects
    if isinstance(st.session_state.get("start_date_1a"), str):
        st.session_state.start_date_1a = datetime.strptime(st.session_state.get("start_date_1a"), "%d-%m-%Y").date()
    if isinstance(st.session_state.get("end_date_1a"), str):
        st.session_state.end_date_1a = datetime.strptime(st.session_state.get("end_date_1a"), "%d-%m-%Y").date()

    st.session_state.start_date_1a = st.date_input("Start Date (less than 2hrs)", value=st.session_state.get("start_date_1a"), min_value=date(1900, 1, 1), max_value=date.today(), help="Choose a start date", format="DD/MM/YYYY")
    st.session_state.end_date_1a = st.date_input("End Date (less than 2hrs)", value=st.session_state.get("end_date_1a"), min_value=date(1900, 1, 1), max_value=date.today(), help="Choose an end date", format="DD/MM/YYYY")

    st.session_state.num_hours_1a = st.number_input("No. of hours (less than 2hrs)", min_value=0, value=st.session_state.get("num_hours_1a", 0))

    # Section 1b: Substantive Numeracy Delivery Course (more than 2 hrs course)
    st.subheader("Section 1b: SUBSTANTIVE NUMERACY DELIVERY COURSE DETAILS (more than 2 hrs)")
    st.session_state.course_code_1b = st.text_input("Course Code (more than 2hrs)", value=st.session_state.get("course_code_1b", ""))
    st.session_state.course_title_1b = st.text_input("Course Title (more than 2hrs)", value=st.session_state.get("course_title_1b", ""))
    st.session_state.delivery_location_1b = st.text_input("Delivery Location (more than 2hrs)", value=st.session_state.get("delivery_location_1b", ""))
    
    # Handle Start Date and End Date as date objects
    if isinstance(st.session_state.get("start_date_1b"), str):
        st.session_state.start_date_1b = datetime.strptime(st.session_state.get("start_date_1b"), "%d-%m-%Y").date()
    if isinstance(st.session_state.get("end_date_1b"), str):
        st.session_state.end_date_1b = datetime.strptime(st.session_state.get("end_date_1b"), "%d-%m-%Y").date()

    st.session_state.start_date_1b = st.date_input("Start Date (more than 2hrs)", value=st.session_state.get("start_date_1b"), min_value=date(1900, 1, 1), max_value=date.today(), help="Choose a start date", format="DD/MM/YYYY")
    st.session_state.end_date_1b = st.date_input("End Date (more than 2hrs)", value=st.session_state.get("end_date_1b"), min_value=date(1900, 1, 1), max_value=date.today(), help="Choose an end date", format="DD/MM/YYYY")

    st.session_state.hours_per_week_1b = st.number_input("Hours per Week (more than 2hrs)", min_value=0, value=st.session_state.get("hours_per_week_1b", 0))
    st.session_state.total_weeks_1b = st.number_input("Total weeks (more than 2hrs)", min_value=0, value=st.session_state.get("total_weeks_1b", 0))
    st.session_state.total_glh_1b = st.number_input("Total GLH (more than 2hrs)", min_value=0, value=st.session_state.get("total_glh_1b", 0))

    # Section 1c: Intervention 5 - Functional Skills Qualification (If applicable)
    st.subheader("Section 1c: INTERVENTION 5: FUNCTIONAL SKILLS QUALIFICATION (If applicable)")
    st.session_state.course_code_1c = st.text_input("Course Code (If applicable)", value=st.session_state.get("course_code_1c", ""))
    st.session_state.course_title_1c = st.text_input("Course Title (If applicable)", value=st.session_state.get("course_title_1c", ""))
    st.session_state.delivery_location_1c = st.text_input("Delivery Location (If applicable)", value=st.session_state.get("delivery_location_1c", ""))
    
    # Handle Start Date and End Date as date objects
    if isinstance(st.session_state.get("start_date_1c"), str):
        st.session_state.start_date_1c = datetime.strptime(st.session_state.get("start_date_1c"), "%d-%m-%Y").date()
    if isinstance(st.session_state.get("end_date_1c"), str):
        st.session_state.end_date_1c = datetime.strptime(st.session_state.get("end_date_1c"), "%d-%m-%Y").date()

    st.session_state.start_date_1c = st.date_input("Start Date (If applicable)", value=st.session_state.get("start_date_1c"), min_value=date(1900, 1, 1), max_value=date.today(), help="Choose a start date", format="DD/MM/YYYY")
    st.session_state.end_date_1c = st.date_input("End Date (If applicable)", value=st.session_state.get("end_date_1c"), min_value=date(1900, 1, 1), max_value=date.today(), help="Choose an end date", format="DD/MM/YYYY")

    st.session_state.hours_per_week_1c = st.number_input("Hours per Week (If applicable)", min_value=0, value=st.session_state.get("hours_per_week_1c", 0))
    st.session_state.total_weeks_1c = st.number_input("Total weeks (If applicable)", min_value=0, value=st.session_state.get("total_weeks_1c", 0))
    st.session_state.total_glh_1c = st.number_input("Total GLH (If applicable)", min_value=0, value=st.session_state.get("total_glh_1c", 0))

    # Next and Back buttons for navigation (in sequence, no columns)
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        # Mandatory check for section 1a and 1b
        if (st.session_state.course_title_1a and
            st.session_state.delivery_location_1a and
            st.session_state.start_date_1a and
            st.session_state.end_date_1a and
            st.session_state.num_hours_1a and
            st.session_state.course_code_1b and
            st.session_state.course_title_1b and
            st.session_state.delivery_location_1b and
            st.session_state.start_date_1b and
            st.session_state.end_date_1b and
            st.session_state.hours_per_week_1b and
            st.session_state.total_weeks_1b and
            st.session_state.total_glh_1b):
            
            # Convert dates to "DD-MM-YYYY" format before proceeding to the next step
            st.session_state.start_date_1a = st.session_state.start_date_1a.strftime("%d-%m-%Y")
            st.session_state.end_date_1a = st.session_state.end_date_1a.strftime("%d-%m-%Y")
            st.session_state.start_date_1b = st.session_state.start_date_1b.strftime("%d-%m-%Y")
            st.session_state.end_date_1b = st.session_state.end_date_1b.strftime("%d-%m-%Y")
            if st.session_state.start_date_1c:
                st.session_state.start_date_1c = st.session_state.start_date_1c.strftime("%d-%m-%Y")
            if st.session_state.end_date_1c:
                st.session_state.end_date_1c = st.session_state.end_date_1c.strftime("%d-%m-%Y")

            st.session_state.step = 3
            st.experimental_rerun()
        else:
            st.warning("Please fill in all mandatory fields in Section 1a and 1b before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 1  # Go back to the previous step
        st.experimental_rerun()

elif st.session_state.step == 3:
    st.title("> Section 2: PERSONAL DETAILS")

    # Title
    if 'title_index' not in st.session_state:
        # Set the default value only if it's not set
        st.session_state.title_index = 0
    
    st.session_state.title = st.selectbox("Title", ["Mr", "Mrs", "Miss", "Ms"], 
                                         index=st.session_state.title_index)

    # Store the title_index so that it persists
    st.session_state.title_index = ["Mr", "Mrs", "Miss", "Ms"].index(st.session_state.title)

    # Legal Sex
    st.session_state.legal_sex = st.radio("Legal Sex", ["Male", "Female"], index=0 if st.session_state.get("legal_sex") == "Male" else 1)

    # Check if dob is a string and convert it back to a date object
    if isinstance(st.session_state.get("dob"), str):
        st.session_state.dob = datetime.strptime(st.session_state.get("dob"), "%d-%m-%Y").date()

    # Date of Birth
    st.session_state.dob = st.date_input(
        label="Date of Birth",  # Label for the field
        value=st.session_state.get("dob"),  # Correctly access dob from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    # National Insurance Number
    st.session_state.ni_number = st.text_input("NI number", value=st.session_state.get("ni_number", ""))

    # Forename
    st.session_state.forename = st.text_input("Forename", value=st.session_state.get("forename", ""))

    # Surname
    st.session_state.surname = st.text_input("Surname", value=st.session_state.get("surname", ""))

    # Previous Surname (Optional)
    st.session_state.previous_surname = st.text_input("Previous Surname (if applicable)", value=st.session_state.get("previous_surname", ""))

    # Current Address
    st.session_state.current_address = st.text_input("Current Address", value=st.session_state.get("current_address", ""))

    # Town
    st.session_state.town = st.text_input("Town", value=st.session_state.get("town", ""))

    # Postcode
    st.session_state.postcode = st.text_input("Postcode", value=st.session_state.get("postcode", ""))

    # Previous Postcodes (Optional, up to 3)
    st.session_state.previous_postcodes = []
    for i in range(3):
        prev_postcode = st.text_input(f"Previous Postcode since 2010 #{i + 1} (if different from current)", value=st.session_state.get(f"previous_postcode_{i+1}", ""))
        if prev_postcode:
            st.session_state.previous_postcodes.append(prev_postcode)

    # Telephone Number
    st.session_state.telephone = st.text_input("Telephone Number", value=st.session_state.get("telephone", ""))

    # Email
    st.session_state.email = st.text_input("Email", value=st.session_state.get("email", "")).strip().replace(" ", "_").lower()

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        # Ensure required fields are filled
        if (is_valid_email(st.session_state.email)):
            if (st.session_state.title and
                st.session_state.legal_sex and
                st.session_state.dob and
                st.session_state.ni_number and
                st.session_state.forename and
                st.session_state.surname and
                st.session_state.current_address and
                st.session_state.town and
                st.session_state.postcode and
                st.session_state.telephone):
                
                # Convert the selected date to the desired string format (DD-MM-YYYY) only when proceeding to the next step
                st.session_state.dob = st.session_state.dob.strftime("%d-%m-%Y")

                st.session_state.step = 4  # Proceed to the next step (Section 3)
                st.experimental_rerun()
            else:
                st.warning("Please fill in all mandatory fields before proceeding.")
        else:
            st.warning("Please enter valid email address.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 2  # Go back to the previous step (Section 1)
        st.experimental_rerun()


elif st.session_state.step == 4:
    st.title("> Section 3: LEARNING RECORDS SERVICE (LRS) PRIVACY NOTICE")

    # Privacy Notice Text
    st.write(
        """
        Please ensure you read the separate Learning Records Service (LRS) Privacy notice given to you along with this form.
        \n* Your previous postcodes will be used to identify the correct learner record on the LRS portal where your information is similar to other learners.
        """
    )

    # Clickable Link to LRS Privacy Notice
    st.markdown(
        '[Find the latest version of the LRS Privacy Notice here.](https://www.gov.uk/government/publications/lrs-privacy-notices/lrs-privacy-notice)',
        unsafe_allow_html=True
    )

    # Radio button for Level 2 or above Maths qualification
    st.session_state.maths_qualification = st.radio(
        "Do you have a Level 2 or above Maths qualification equivalent to GCSE at grade C or 4 and above or functional skills level 2?",
        ["Yes", "No"],
        index=0 if st.session_state.get("maths_qualification") == "Yes" else 1
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        st.session_state.step = 5  # Proceed to the next step (Section 4)
        st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 3  # Go back to the previous step (Section 2)
        st.experimental_rerun()

elif st.session_state.step == 5:
    st.title("> Section 4: RESIDENCY ELIGIBILITY")

    # Question 1: Living in the UK/EEA for at least 3 years
    st.session_state.uk_eea_residency = st.radio(
        "Have you been living in the UK or an EEA country continuously for at least 3 years before your learning start date?",
        ["Yes", "No"],
        index=0 if st.session_state.get("uk_eea_residency") == "Yes" else 1
    )

    # Question 2: Nationality or EEA settlement status
    st.session_state.nationality_status = st.radio(
        "Are you one of the following: A UK national? OR An Irish national? An EEA national with pre-settled or settled status under the EU Settlement Scheme?",
        ["Yes", "No"],
        index=0 if st.session_state.get("nationality_status") == "Yes" else 1
    )

    # If 'No' is selected for either question, provide instructions for residency evidence
    if st.session_state.uk_eea_residency == "No" or st.session_state.nationality_status == "No":
        st.warning("You will need to show us evidence of your residency status. Please ensure that the evidence is available for verification.")

        # Partner Note for immigration documents check
        st.info(
            """
            Partner: If the learner has answered ‘No’ to either of the 2 residency questions above, you will be required to check the learner’s immigration documents to determine if they meet the residency eligibility requirements of the Multiply project.
            \nYou must document the detail of the learner’s immigration documents on an appendix B1 form.
            \nCopies of learner immigration documents must not be taken.
            """
        )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        st.session_state.step = 6  # Proceed to the next step (Section 5)
        st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 4  # Go back to the previous step (Section 3)
        st.experimental_rerun()

elif st.session_state.step == 6:
    st.title("> Section 5: ETHNICITY")

    # Broad Ethnicity Category Selection
    broad_ethnicity_options = [
        "Select a category",  # Placeholder option
        "White",
        "Mixed/Multiple Ethnic Group",
        "Asian/Asian British",
        "Black/African/Caribbean/Black British",
        "Other Ethnic group"
    ]
    
    # Detect if the broad ethnicity category has changed
    if 'prev_ethnicity_category' in st.session_state:
        if st.session_state.ethnicity_category != st.session_state.prev_ethnicity_category:
            # Reset specific ethnicity detail when broad category changes
            st.session_state.ethnicity_detail = None

    st.session_state.ethnicity_category = st.selectbox(
        "Please choose your broad ethnicity category:",
        broad_ethnicity_options,
        index=broad_ethnicity_options.index(st.session_state.get("ethnicity_category", "Select a category"))
    )

    # Store the current category to check for changes
    st.session_state.prev_ethnicity_category = st.session_state.ethnicity_category

    # Define options based on broad category
    if st.session_state.ethnicity_category == "White":
        ethnicity_detail_options = [
            "Select an option",  # Placeholder option
            "31 English/Welsh/Scottish/Northern Irish/British",
            "32 Irish",
            "33 Gypsy or Irish Traveller",
            "34 Any other White background"
        ]
    elif st.session_state.ethnicity_category == "Mixed/Multiple Ethnic Group":
        ethnicity_detail_options = [
            "Select an option",  # Placeholder option
            "35 White and Black Caribbean",
            "36 White and Black African",
            "37 White and Asian",
            "38 Any other Mixed/multiple ethnic background"
        ]
    elif st.session_state.ethnicity_category == "Asian/Asian British":
        ethnicity_detail_options = [
            "Select an option",  # Placeholder option
            "39 Indian",
            "40 Pakistani",
            "41 Bangladeshi",
            "42 Chinese",
            "43 Any other Asian background"
        ]
    elif st.session_state.ethnicity_category == "Black/African/Caribbean/Black British":
        ethnicity_detail_options = [
            "Select an option",  # Placeholder option
            "44 African",
            "45 Caribbean",
            "46 Any other Black/African/Caribbean background"
        ]
    elif st.session_state.ethnicity_category == "Other Ethnic group":
        ethnicity_detail_options = [
            "Select an option",  # Placeholder option
            "47 Arab",
            "98 Any other ethnic group"
        ]
    else:
        ethnicity_detail_options = []

    # Specific Ethnicity Detail Selection
    st.session_state.ethnicity_detail = st.selectbox(
        "Please select your specific ethnicity:",
        ethnicity_detail_options,
        index=ethnicity_detail_options.index(st.session_state.get("ethnicity_detail", "Select an option")) if st.session_state.get("ethnicity_detail") in ethnicity_detail_options else 0
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click with validation
    if next_clicked:
        if st.session_state.ethnicity_category == "Select a category":
            st.warning("Please select a valid broad ethnicity category.")
        elif st.session_state.ethnicity_detail == "Select an option" or not st.session_state.ethnicity_detail:
            st.warning("Please select a valid specific ethnicity.")
        else:
            st.session_state.step = 7  # Proceed to the next step (Section 6)
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 5  # Go back to the previous step (Section 4)
        st.experimental_rerun()

elif st.session_state.step == 7:
    st.title("> Section 6: LEARNING DIFFICULTIES, DISABILITIES AND HEALTH PROBLEMS")

    # Initial Question (Check for existing value in st.session_state)
    # Use the return value of st.radio() instead of modifying st.session_state directly
    has_difficulties = st.radio(
        "Do you have any learning difficulties, disabilities and/or health problems?",
        ["No", "Yes"],
        index=0 if st.session_state.get("has_difficulties", "No") == "No" else 1
    )
    st.session_state.has_difficulties = has_difficulties  # Update st.session_state based on radio input

    if st.session_state.has_difficulties == "Yes":
        # List of Options
        options = [
            "4 Vision Impairment",
            "5 Hearing Impairment",
            "6 Disability affecting mobility",
            "7 Profound complex disabilities",
            "8 Social and emotional difficulties",
            "9 Mental health difficulty",
            "10 Moderate learning difficulty",
            "11 Severe learning difficulty",
            "12 Dyslexia",
            "13 Dyscalculia",
            "14 Autism spectrum disorder",
            "15 Aspergers syndrome",
            "16 Temporary disability after illness/accident (i.e post viral)",
            "93 Other physical disability",
            "94 Other specific learning difficulty (e.g. Dyspraxia)",
            "95 Other medical condition (e.g epilepsy, asthma diabetes)",
            "96 Other learning difficulty",
            "97 Other disability",
            "98 Prefer not to say"
        ]

        # Multi-select for relevant options: Only update selected_difficulties if the user has interacted with the widget
        selected_difficulties = st.multiselect(
            "Please choose all those that affect you: (Caution: You may need to click every option twice to get it reflected.)",
            options,
            default=st.session_state.get("selected_difficulties", [])
        )

        # Only update session state if the user has interacted with the multi-select
        if selected_difficulties != st.session_state.get("selected_difficulties"):
            st.session_state.selected_difficulties = selected_difficulties

        # Additional Input if multiple options are selected
        if len(st.session_state.selected_difficulties) > 1:
            # Ensure that the most_affecting option is part of the selected difficulties
            if st.session_state.get("most_affecting") not in st.session_state.selected_difficulties:
                st.session_state.most_affecting = st.session_state.selected_difficulties[0]  # Default to the first option

            # Only show the options selected in the multi-select
            st.session_state.most_affecting = st.selectbox(
                "If you have marked more than one box, please record the code number of the condition that affects you the most:",
                st.session_state.selected_difficulties,
                index=st.session_state.selected_difficulties.index(st.session_state.get("most_affecting"))
            )
        else:
            # Reset the most affecting condition if there's only one selected
            st.session_state.most_affecting = None
    else:
        # Reset the selected difficulties and most affecting if No is selected
        st.session_state.selected_difficulties = []
        st.session_state.most_affecting = None

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if st.session_state.has_difficulties == "Yes" and not st.session_state.selected_difficulties:
            st.warning("Please select at least one condition that affects you.")
        else:
            st.session_state.step = 8  # Proceed to the next step (Section 7)
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 6  # Go back to the previous step (Section 5)
        st.experimental_rerun()

elif st.session_state.step == 8:
    st.title("> Section 7: HIGHEST LEVEL OF QUALIFICATION")

    # List of qualification options
    qualifications = [
        "No qualifications",
        "Entry level / qualification below level 1",
        "Level 1",
        "Level 2",
        "Full level 2",
        "Level 3",
        "Full level 3",
        "Level 4",
        "Level 5",
        "Level 6",
        "Level 7 and above",
        "Other qualification, level not known"
    ]

    # Qualification selectbox
    st.session_state.qualification = st.selectbox(
        "Please indicate the highest level of qualification you currently hold:",
        qualifications,
        index=qualifications.index(st.session_state.get("qualification", "No qualifications"))
    )

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if st.session_state.qualification:
            st.session_state.step = 9  # Proceed to the next step
            st.experimental_rerun()
        else:
            st.warning("Please select your highest qualification before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 7  # Go back to the previous step (Section 6)
        st.experimental_rerun()

elif st.session_state.step == 9:
    st.title("> Section 8: EMPLOYMENT & BENEFITS")

    # Employment Status Subheader
    st.subheader("Please confirm your current employment status details:")

    # Employment status options
    employment_status_options = [
        "I am in paid employment [10]",
        "I am self employed [10 & SEI 1]",
        "I am not in paid employment but looking for work and available to start work [11]",
        "I am not in paid employment, not looking for work and/or not available to start work [12]",
        "I am retired [80]",
        "I am in full time education or training [PEI 1]"
    ]

    # Get the current employment status or set a valid default
    employment_status = st.session_state.get("employment_status", employment_status_options[0])

    # Ensure the employment status is valid before using it in the index function
    if employment_status not in employment_status_options:
        employment_status = employment_status_options[0]

    # Employment status radio button
    st.session_state.employment_status = st.radio(
        "Employment Status",
        employment_status_options,
        index=employment_status_options.index(employment_status)
    )

    # Working hours options if the user is in paid employment
    working_hours_options = [
        "0 to 10 hours per week [EII5]",
        "11 to 20 hours per week [EII6]",
        "21 to 30 hours per week [EII7]",
        "31 hours per week or more [EII8]"
    ]

    # Unemployment duration options if the user is not in paid employment
    unemployment_duration_options = [
        "Less than 6 months [LOU01]",
        "6 – 11 months [LOU02]",
        "12 – 23 months [LOU03]",
        "24 – 35 months [LOU04]",
        "36 months or more [LOU05]",
        "I have been made redundant [OET1]"
    ]

    # Conditional sub-options based on employment status
    if st.session_state.employment_status == "I am in paid employment [10]":
        # Get the current working hours or set a valid default
        working_hours = st.session_state.get("working_hours", working_hours_options[0])

        # Ensure the working hours are valid before using them in the index function
        if working_hours not in working_hours_options or working_hours == '':
            working_hours = working_hours_options[0]

        # Working hours select box
        st.session_state.working_hours = st.selectbox(
            "And I work for:",
            working_hours_options,
            index=working_hours_options.index(working_hours)  # Using validated working_hours
        )

    elif st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]":
        # Get the current unemployment duration or set a valid default
        unemployment_duration = st.session_state.get("unemployment_duration", unemployment_duration_options[0])

        # Ensure the unemployment duration is valid before using it in the index function
        if unemployment_duration not in unemployment_duration_options or unemployment_duration == '':
            unemployment_duration = unemployment_duration_options[0]

        # Unemployment duration select box
        st.session_state.unemployment_duration = st.selectbox(
            "I have been unemployed for:",
            unemployment_duration_options,
            index=unemployment_duration_options.index(unemployment_duration)  # Using validated unemployment_duration
        )



    # Benefit Status Subheader
    st.subheader("Please indicate if you are currently in receipt of one of the following benefits:")

    # Benefit options as checkboxes
    st.session_state.job_seekers_allowance = st.checkbox("Job seekers allowance [BSI1]", value=st.session_state.get("job_seekers_allowance", False))
    st.session_state.esa = st.checkbox("Employment & Support Allowance (ESA) [BSI5]", value=st.session_state.get("esa", False))
    st.session_state.universal_credit = st.checkbox("Universal credit [BSI4]", value=st.session_state.get("universal_credit", False))
    st.session_state.other_benefit = st.checkbox("Other state benefit [BSI6]", value=st.session_state.get("other_benefit", False))

    # If any benefit is selected, prompt for National Insurance number
    if st.session_state.job_seekers_allowance or st.session_state.esa or st.session_state.universal_credit or st.session_state.other_benefit:
        st.session_state.ni_number = st.text_input(
            "If you have ticked one of these benefit boxes, please provide your National Insurance number:",
            value=st.session_state.get("ni_number", ""),
            max_chars=9,
            help="(2 letters, 6 numbers, 1 letter)"
        )


    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.job_seekers_allowance or st.session_state.esa or st.session_state.universal_credit or st.session_state.other_benefit) and not st.session_state.ni_number:
            st.warning("Please provide your National Insurance number if you are receiving any benefits.")
        else:
            st.session_state.step = 10  # Proceed to the next step
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 8  # Go back to the previous step
        st.experimental_rerun()

elif st.session_state.step == 10:
    st.title("> Section 9: EDUCATION AND SKILLS FUNDING AGENCY (ESFA) PRIVACY NOTICE")

    # Privacy Notice Content
    st.write("""
    Please ensure you read the separate Education and Skills Funding Agency (ESFA) privacy notice given to you along with this form. 
    You may also find the latest version of the ESFA privacy notice on their website:
    [Privacy Notice for Key Stage 5 and Adult Education](https://www.gov.uk/government/publications/privacy-notice-for-key-stage-5-and-adult-education/privacy-notice-for-key-stage-5-and-adult-education)
    """)

    # Checkbox to confirm reading the privacy notice
    st.session_state.privacy_agreed = st.checkbox(
        "I have read and agree to the ESFA privacy notice",
        value=st.session_state.get("privacy_agreed", False)
    )

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if st.session_state.privacy_agreed:
            st.session_state.step = 11  # Proceed to the next step
            st.experimental_rerun()
        else:
            st.warning("Please read and agree to the ESFA privacy notice before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 9  # Go back to the previous step
        st.experimental_rerun()

elif st.session_state.step == 11:
    st.title("> Section 10: CLS Marketing, Research and Profiling Permissions")

    # Permissions for contact about courses, offers, research, etc.
    st.subheader("Please choose the boxes if you wish to be contacted by us in the future:")
    
    st.session_state.marketing_courses_offers = st.checkbox(
        "About courses, learning opportunities and relevant promotional offers",
        value=st.session_state.get("marketing_courses_offers", False)
    )
    
    st.session_state.research_profiling = st.checkbox(
        "For research, profiling, and quality improvement purposes",
        value=st.session_state.get("research_profiling", False)
    )

    # Contact preferences
    st.subheader("Please indicate your CLS contact preferences")
    
    st.session_state.contact_email = st.checkbox(
        "Email", 
        value=st.session_state.get("contact_email", False)
    )
    
    st.session_state.contact_post = st.checkbox(
        "Post", 
        value=st.session_state.get("contact_post", False)
    )
    
    st.session_state.contact_phone = st.checkbox(
        "Phone", 
        value=st.session_state.get("contact_phone", False)
    )
    
    st.session_state.contact_text = st.checkbox(
        "Text", 
        value=st.session_state.get("contact_text", False)
    )

    # CLS Privacy Notice link
    st.write("The CLS privacy notice can be found on our website at: [CLS Privacy Notice](https://previstaltd.sharepoint.com/:b:/g/allcompanydrive/EaeINtmtm1FOiofrEEmUUw0BoqHaocif2spsWUjjY68nKg?e=mClImT)")

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        # At least one contact preference should be selected if user opted for marketing or research
        if (st.session_state.marketing_courses_offers or st.session_state.research_profiling) and not (
            st.session_state.contact_email or 
            st.session_state.contact_post or 
            st.session_state.contact_phone or 
            st.session_state.contact_text):
            st.warning("Please select at least one contact preference if you wish to be contacted.")
        else:
            st.session_state.step = 12  # Proceed to the next step
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 10  # Go back to the previous step
        st.experimental_rerun()


elif st.session_state.step == 12:
    st.title("> Section 11: LEARNER DECLARATION")

    # Declaration checkboxes
    st.session_state.confirm_info_correct = st.checkbox(
        "I confirm that all the information I have provided on this form is correct.",
        value=st.session_state.get("confirm_info_correct", False)
    )
    
    st.session_state.agree_inform_changes = st.checkbox(
        "I agree to inform the Multiply Programme of any changes to this information during my learning.",
        value=st.session_state.get("agree_inform_changes", False)
    )

    st.session_state.confirm_read_privacy = st.checkbox(
        "I confirm I have read the LRS and ESFA privacy notices.",
        value=st.session_state.get("confirm_read_privacy", False)
    )

    # Signature Box - Using streamlit_drawable_canvas for a signature box simulation
    st.subheader("Signature:")
    st.write("Please draw your signature below:")

    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="white",  # White background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="canvas",
    )
    st.session_state.signature = canvas_result.image_data

    # Set today's date automatically and display it
    st.session_state.signature_date = date.today().strftime("%d-%m-%Y")
    st.write(f"Date: **{st.session_state.signature_date}**")

    # Submit button
    submit_clicked = st.button("Submit")

###############################

    # Handle Submit button click
    if submit_clicked:
        # Check if all checkboxes are checked and signature is drawn
        if (st.session_state.confirm_info_correct and
            st.session_state.agree_inform_changes and
            st.session_state.confirm_read_privacy and
            is_signature_drawn(st.session_state.signature)):
            st.warning('Please wait! We are currently processing. . . .', icon="🚨")

            with st.spinner('Wait for it...'):
                # FILL TEMPLATE:
                st.session_state.placeholder_values = {
                    # Section 1a: Engagement Session Details (less than 2hrs)
                    'ph1': st.session_state.course_title_1a,
                    'ph2': st.session_state.delivery_location_1a,
                    'ph3': st.session_state.start_date_1a,
                    'ph4': st.session_state.end_date_1a,
                    'ph5': st.session_state.num_hours_1a,

                    # Section 1b: Substantive Numeracy Delivery Course Details (more than 2 hrs)
                    'ph6': st.session_state.course_code_1b,
                    'ph7': st.session_state.course_title_1b,
                    'ph8': st.session_state.delivery_location_1b,
                    'ph9': st.session_state.start_date_1b,
                    'ph10': st.session_state.end_date_1b,
                    'ph11': st.session_state.hours_per_week_1b,
                    'ph12': st.session_state.total_weeks_1b,
                    'ph13': st.session_state.total_glh_1b,

                    # Section 1c: Functional Skills Qualification (If applicable)
                    'ph14': st.session_state.course_code_1c,
                    'ph15': st.session_state.course_title_1c,
                    'ph16': st.session_state.delivery_location_1c,
                    'ph17': st.session_state.start_date_1c,
                    'ph18': st.session_state.end_date_1c,
                    'ph19': st.session_state.hours_per_week_1c,
                    'ph19a': st.session_state.total_weeks_1c,
                    'ph20': st.session_state.total_glh_1c,

                    # Section 2: Personal Details
                    'ph21': st.session_state.title,
                    'ph22': "X" if st.session_state.legal_sex == "Male" else "",
                    'ph23': "X" if st.session_state.legal_sex == "Female" else "",
                    'ph24': st.session_state.dob,
                    'ph25': st.session_state.ni_number,
                    'ph26': st.session_state.forename,
                    'ph27': st.session_state.surname,
                    'ph28': st.session_state.previous_surname,
                    'ph29': st.session_state.current_address,
                    'ph30': st.session_state.town,
                    'ph31': st.session_state.postcode,
                    'ph32': st.session_state.previous_postcodes[0] if len(st.session_state.previous_postcodes) > 0 else "",
                    'ph33': st.session_state.previous_postcodes[1] if len(st.session_state.previous_postcodes) > 1 else "",
                    'ph34': st.session_state.previous_postcodes[2] if len(st.session_state.previous_postcodes) > 2 else "",
                    'ph35': st.session_state.telephone,
                    'ph36': st.session_state.email,

                    # Section 3: LRS Privacy Notice
                    'ph37': "X" if st.session_state.maths_qualification == "Yes" else "",
                    'ph38': "X" if st.session_state.maths_qualification == "No" else "",

                    # Section 4: Residency Eligibility
                    'ph39': "X" if st.session_state.uk_eea_residency == "Yes" else "",
                    'ph40': "X" if st.session_state.uk_eea_residency == "No" else "",
                    'ph41': "X" if st.session_state.nationality_status == "Yes" else "",
                    'ph42': "X" if st.session_state.nationality_status == "No" else "",

                    # Section 5: Ethnicity
                    'ph43': "X" if st.session_state.ethnicity_detail == "31 English/Welsh/Scottish/Northern Irish/British" else "",
                    'ph44': "X" if st.session_state.ethnicity_detail == "32 Irish" else "",
                    'ph45': "X" if st.session_state.ethnicity_detail == "33 Gypsy or Irish Traveller" else "",
                    'ph46': "X" if st.session_state.ethnicity_detail == "34 Any other White background" else "",
                    'ph47': "X" if st.session_state.ethnicity_detail == "35 White and Black Caribbean" else "",
                    'ph48': "X" if st.session_state.ethnicity_detail == "36 White and Black African" else "",
                    'ph49': "X" if st.session_state.ethnicity_detail == "37 White and Asian" else "",
                    'ph50': "X" if st.session_state.ethnicity_detail == "38 Any other Mixed/multiple ethnic background" else "",
                    'ph51': "X" if st.session_state.ethnicity_detail == "39 Indian" else "",
                    'ph52': "X" if st.session_state.ethnicity_detail == "40 Pakistani" else "",
                    'ph53': "X" if st.session_state.ethnicity_detail == "41 Bangladeshi" else "",
                    'ph54': "X" if st.session_state.ethnicity_detail == "42 Chinese" else "",
                    'ph55': "X" if st.session_state.ethnicity_detail == "43 Any other Asian background" else "",
                    'ph56': "X" if st.session_state.ethnicity_detail == "44 African" else "",
                    'ph57': "X" if st.session_state.ethnicity_detail == "45 Caribbean" else "",
                    'ph58': "X" if st.session_state.ethnicity_detail == "46 Any other Black/African/Caribbean background" else "",
                    'ph59': "X" if st.session_state.ethnicity_detail == "47 Arab" else "",
                    'ph60': "X" if st.session_state.ethnicity_detail == "98 Any other ethnic group" else "",

                    # Section 6: Learning Difficulties and Health Problems
                    'ph61': "X" if st.session_state.has_difficulties == "Yes" else "",
                    'ph62': "X" if st.session_state.has_difficulties == "No" else "",
                    'ph63': "X" if "4 Vision Impairment" in st.session_state.selected_difficulties else "",
                    'ph64': "X" if "5 Hearing Impairment" in st.session_state.selected_difficulties else "",
                    'ph65': "X" if "6 Disability affecting mobility" in st.session_state.selected_difficulties else "",
                    'ph66': "X" if "7 Profound complex disabilities" in st.session_state.selected_difficulties else "",
                    'ph67': "X" if "8 Social and emotional difficulties" in st.session_state.selected_difficulties else "",
                    'ph68': "X" if "9 Mental health difficulty" in st.session_state.selected_difficulties else "",
                    'ph69': "X" if "10 Moderate learning difficulty" in st.session_state.selected_difficulties else "",
                    'ph70': "X" if "11 Severe learning difficulty" in st.session_state.selected_difficulties else "",
                    'ph71': "X" if "12 Dyslexia" in st.session_state.selected_difficulties else "",
                    'ph72': "X" if "13 Dyscalculia" in st.session_state.selected_difficulties else "",
                    'ph73': "X" if "14 Autism spectrum disorder" in st.session_state.selected_difficulties else "",
                    'ph74': "X" if "15 Aspergers syndrome" in st.session_state.selected_difficulties else "",
                    'ph75': "X" if "16 Temporary disability after illness/accident (i.e post viral)" in st.session_state.selected_difficulties else "",
                    'ph76': "X" if "93 Other physical disability" in st.session_state.selected_difficulties else "",
                    'ph77': "X" if "94 Other specific learning difficulty (e.g. Dyspraxia)" in st.session_state.selected_difficulties else "",
                    'ph78': "X" if "95 Other medical condition (e.g epilepsy, asthma diabetes)" in st.session_state.selected_difficulties else "",
                    'ph79': "X" if "96 Other learning difficulty" in st.session_state.selected_difficulties else "",
                    'ph80': "X" if "97 Other disability" in st.session_state.selected_difficulties else "",
                    'ph81': "X" if "98 Prefer not to say" in st.session_state.selected_difficulties else "",
                    'ph82': st.session_state.most_affecting.split(" ")[0] if st.session_state.most_affecting else "",  # Only the code number


                    # Section 7: Highest Level of Qualification
                    'ph83': "X" if st.session_state.qualification == "No qualifications" else "",
                    'ph84': "X" if st.session_state.qualification == "Entry level / qualification below level 1" else "",
                    'ph85': "X" if st.session_state.qualification == "Level 1" else "",
                    'ph86': "X" if st.session_state.qualification == "Level 2" else "",
                    'ph87': "X" if st.session_state.qualification == "Full level 2" else "",
                    'ph88': "X" if st.session_state.qualification == "Level 3" else "",
                    'ph89': "X" if st.session_state.qualification == "Full level 3" else "",
                    'ph90': "X" if st.session_state.qualification == "Level 4" else "",
                    'ph91': "X" if st.session_state.qualification == "Level 5" else "",
                    'ph92': "X" if st.session_state.qualification == "Level 6" else "",
                    'ph93': "X" if st.session_state.qualification == "Level 7 and above" else "",
                    'ph94': "X" if st.session_state.qualification == "Other qualification, level not known" else "",


                    # Section 8: Employment & Benefit
                    'ph95': "X" if st.session_state.employment_status == "I am in paid employment [10]" else "",
                    # Update placeholder logic for working hours status
                    'ph96' : "X" if (st.session_state.employment_status == "I am in paid employment [10]" and st.session_state.working_hours == "0 to 10 hours per week [EII5]") else "",
                    'ph97' : "X" if (st.session_state.employment_status == "I am in paid employment [10]" and st.session_state.working_hours == "11 to 20 hours per week [EII6]") else "",
                    'ph98' : "X" if (st.session_state.employment_status == "I am in paid employment [10]" and st.session_state.working_hours == "21 to 30 hours per week [EII7]") else "",
                    'ph99' : "X" if (st.session_state.employment_status == "I am in paid employment [10]" and st.session_state.working_hours == "31 hours per week or more [EII8]") else "",
                    'ph100': "X" if st.session_state.employment_status == "I am self employed [10 & SEI 1]" else "",
                    'ph101': "X" if st.session_state.employment_status == "I am not in paid employment but looking for work and available to start work [11]" else "",
                    'ph102': "X" if st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]" else "",
                    # Update placeholder logic for unemployment status
                    'ph103' : "X" if (st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]" and st.session_state.unemployment_duration == "Less than 6 months [LOU01]") else "",
                    'ph104' : "X" if (st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]" and st.session_state.unemployment_duration == "6 – 11 months [LOU02]") else "",
                    'ph105' : "X" if (st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]" and st.session_state.unemployment_duration == "12 – 23 months [LOU03]") else "",
                    'ph106' : "X" if (st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]" and st.session_state.unemployment_duration == "24 – 35 months [LOU04]") else "",
                    'ph107' : "X" if (st.session_state.employment_status == "I am not in paid employment, not looking for work and/or not available to start work [12]" and st.session_state.unemployment_duration == "36 months or more [LOU05]") else "",

                    'ph108': "X" if st.session_state.unemployment_duration == "I have been made redundant [OET1]" else "",
                    'ph109': "X" if st.session_state.employment_status == "I am retired [80]" else "",
                    'ph110': "X" if st.session_state.employment_status == "I am in full time education or training [PEI 1]" else "",

                    # Benefits
                    'ph111': "X" if st.session_state.job_seekers_allowance else "",
                    'ph112': "X" if st.session_state.esa else "",
                    'ph113': "X" if st.session_state.universal_credit else "",
                    'ph114': "X" if st.session_state.other_benefit else "",

                    # National Insurance number (only if any benefit is selected)
                    'ph115': st.session_state.ni_number if (st.session_state.job_seekers_allowance or st.session_state.esa or st.session_state.universal_credit or st.session_state.other_benefit) else "",


                    # Section 10: CLS Marketing and Permissions

                    # Marketing Permissions
                    'ph116': "X" if st.session_state.marketing_courses_offers else "",
                    'ph117': "X" if st.session_state.research_profiling else "",

                    # Contact Preferences
                    'ph118': "X" if st.session_state.contact_email else "",
                    'ph119': "X" if st.session_state.contact_phone else "",
                    'ph120': "X" if st.session_state.contact_post else "",
                    'ph121': "X" if st.session_state.contact_text else "",


                    # Section 11: Learner Declaration
                    'ph122': st.session_state.signature_date
                }

                # Remove leading/trailing spaces, then replace internal spaces with underscores, and convert to lowercase
                safe_first_name = st.session_state.forename.strip().replace(" ", "_").lower()
                safe_family_name = st.session_state.surname.strip().replace(" ", "_").lower()

                # Define input and output paths
                template_file = "resources/ph_multiply_surrey.docx"
                modified_file = f"MultiplySurrey_Form_Submission_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.docx"

                signature_path = f'signature_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.png'            
                resized_image_path = f'resized_signature_image_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.png'

                try:
                    signature_image = PILImage.fromarray(
                        st.session_state.signature.astype('uint8'), 'RGBA')
                    signature_image.save(signature_path)

                    # Open and resize the image
                    print(f"Opening image file: {signature_path}")
                    resized_image = PILImage.open(signature_path)
                    print(f"Original image size: {resized_image.size}")
                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                    resized_image.save(resized_image_path)  # Save resized image to a file
                    print(f"Resized image saved to: {resized_image_path}")

                    replace_placeholders(template_file, modified_file, st.session_state.placeholder_values, resized_image_path)
                except Exception as e:
                    # Display the error message on the screen
                    st.error('Please wait, form will reprocess and will give you the option again to submit in 10 SECONDS automatically')
                    st.error(f"Please take screenshot of the following error and share with Developer: \n{str(e)}")
                    time.sleep(12)

                    st.session_state.submission_done = False
                    st.session_state.step = 12
                    st.experimental_rerun()


                # Email
                # Sender email credentials
                # Credentials: Streamlit host st.secrets
                # sender_email = 'dummy'
                # sender_password = 'dummy'
                sender_email = st.secrets["sender_email"]
                sender_password = st.secrets["sender_password"]

                # Credentials: Local env
                # load_dotenv()                                     # uncomment import of this library!
                # sender_email = os.getenv('EMAIL')
                # sender_password = os.getenv('PASSWORD')
                team_email = [sender_email]
                # team_email = ['muhammadoa@prevista.co.uk']
                # receiver_email = sender_email
                # receiver_email = 'muhammadoa@prevista.co.uk'

                # learner_email = [st.session_state.email]
                
                # subject_team = f"MultiplySurrey: {st.session_state.selected_option} {st.session_state.hear_about}_{st.session_state.hother_source}_{st.session_state.forename}_{st.session_state.surname} Submission Date: {date.today()}"
                subject_team = f"MultiplySurrey: {st.session_state.forename}_{st.session_state.surname} Submission Date: {date.today()}"
                body_team = f'''Prevista Multiply Surrey Form submitted. Please find attached file.'''

                # subject_learner = "Thank You for Your Interest in The Skills Bootcamp!"
                body_learner = f"""
                <html>
                <body>
                    <p>Dear {st.session_state.forename} {st.session_state.surname},</p>

                    <p>Thank you for expressing your interest in Bootcamp at PREVISTA. We are excited to guide you through the next steps of the enrollment process.</p>

                    <p><strong>What’s Next?</strong></p>
                    <ol>
                        <li><strong>Enrollment Communication:</strong> One of our representatives will be contacting you within the next few days to complete your enrollment. Please keep an eye out for our message to finalize your registration details.</li>
                        <li><strong>Course Start Date:</strong> Once your enrollment is confirmed, we will send you the schedule for the course start date.</li>
                        <li><strong>Orientation Session:</strong> You will be invited to an orientation session where you can learn more about the platform, meet your instructors, and connect with other learners.</li>
                    </ol>

                    <p>If you have any immediate questions, please feel free to reach out to us at PrevistaAdmissions@prevista.co.uk.</p>

                    <p>We look forward to speaking with you soon and welcoming you to our learning community!</p>

                    <p>Best regards,</p>
                    <p>Student Admissions Team<br>
                    PREVISTA<br>
                    PREPARING YOU TODAY FOR OPPORTUNITIES OF TOMORROW</p>
                </body>
                </html>
                """

                # Local file path
                local_file_path = modified_file

                # Send email to team with attachments
                if st.session_state.files or local_file_path:
                    send_email_with_attachments(sender_email, sender_password, team_email, subject_team, body_team, st.session_state.files, local_file_path)                
                    st.session_state.submission_done = True
                # Send thank you email to learner
                # send_email_with_attachments(sender_email, sender_password, learner_email, subject_learner, body_learner)

                # st.success("Processing Complete!")
                # st.write("Someone will get in touch with you soon.")
###############################
        else:
            st.warning("Please ensure all declarations are agreed to, and the signature is drawn.")

    # Add a warning before the back button
    st.info("If you go back, you will have to re-sign the form.")

    # Navigation buttons
    back_clicked = st.button("Back", disabled=st.session_state.submission_done)

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 11  # Go back to the previous step
        st.experimental_rerun()

if st.session_state.submission_done:
    st.success("Submission Finished!")
    try:
        # file download button
        with open(modified_file, 'rb') as f:
            file_contents = f.read()
            st.download_button(
                label="Download Your Response",
                data=file_contents,
                file_name=modified_file,
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        # clear session state
        st.session_state.files = []
        last()
        st.write("Please close the form. After submission you can't go back!")
        st.snow()

    except Exception as e:
        st.write("Unable to download the file. Please whatsapp learner name to +447405327072 for verificatino of submission.")
        st.error('Please wait, form will reprocess and will give you the option again to submit in 10 SECONDS')
        time.sleep(12)

        st.session_state.submission_done = False
        st.session_state.step = 12
        st.experimental_rerun()

# streamlit run app.py
# Dev : https://linkedin.com/in/osamatech786